#!/usr/bin/env python3
"""DOCX exporter for chapter 6/7 report."""

from __future__ import annotations

from pathlib import Path
from typing import Any


def export_docx(content: dict[str, Any], output_path: Path) -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception as exc:  # pragma: no cover - dependency gate
        raise RuntimeError("python-docx is required. Install dependency: python-docx") from exc

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    title = str(content.get("title", "六七章自动生成报告"))
    doc.add_heading(title, level=0)

    outline_path = str(content.get("outline_path", ""))
    if content.get("outline_exists"):
        doc.add_paragraph(f"模板来源：{outline_path}")
    else:
        doc.add_paragraph(f"模板来源缺失：{outline_path}（已按内置模板生成）")

    for block in content.get("blocks", []):
        btype = block.get("type")
        if btype == "heading":
            lv = int(block.get("level", 1))
            lv = max(1, min(4, lv))
            doc.add_heading(str(block.get("text", "")), level=lv)
        elif btype == "paragraph":
            doc.add_paragraph(str(block.get("text", "")))
        elif btype == "bullet":
            doc.add_paragraph(str(block.get("text", "")), style="List Bullet")
        elif btype == "image":
            caption = str(block.get("caption", ""))
            path = Path(str(block.get("path", "")))
            exists = bool(block.get("exists", False)) and path.exists()
            if exists:
                doc.add_picture(str(path), width=Inches(6.0))
                p = doc.add_paragraph(caption)
                if p.runs:
                    p.runs[0].italic = True
            else:
                doc.add_paragraph(f"图像缺失：{caption}（{path.as_posix()}）")
        elif btype == "table":
            caption = str(block.get("caption", ""))
            headers = [str(h) for h in block.get("headers", [])]
            rows = [[str(c) for c in row] for row in block.get("rows", [])]
            if caption:
                p = doc.add_paragraph(caption)
                if p.runs:
                    p.runs[0].italic = True
            if headers:
                tb = doc.add_table(rows=1 + len(rows), cols=len(headers))
                tb.style = "Table Grid"
                for j, h in enumerate(headers):
                    tb.cell(0, j).text = h
                for i, row in enumerate(rows, start=1):
                    for j in range(len(headers)):
                        tb.cell(i, j).text = row[j] if j < len(row) else ""

    evidence_rows = content.get("evidence_rows", [])
    if evidence_rows:
        doc.add_heading("证据附录", level=1)
        seen: set[tuple[str, str, str]] = set()
        for row in evidence_rows:
            section = str(row.get("section", ""))
            statement = str(row.get("statement", ""))
            path = str(row.get("evidence_path", ""))
            key = (section, statement, path)
            if key in seen:
                continue
            seen.add(key)
            doc.add_paragraph(f"[{section}] {statement} -> {path}", style="List Bullet")

    doc.save(str(output_path))
